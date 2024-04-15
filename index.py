import pdfplumber
import docx
import re
from openpyxl import Workbook, load_workbook
import os
from win32com.client import Dispatch


def list_files_in_directory(directory_path):
    # Check if the directory path exists
    if not os.path.exists(directory_path):
        print(f"Directory '{directory_path}' does not exist.")
        return
    
    # Check if the directory path is a directory
    if not os.path.isdir(directory_path):
        print(f"'{directory_path}' is not a directory.")
        return
    
    # Get the list of files in the directory
    files = [f for f in os.listdir(directory_path) if os.path.isfile(os.path.join(directory_path, f))]
    
    # Print the list of files
    # if files:
    #     return files
    # else:
    #     print("No files found in the directory.")
    
    return files if files else None


def convert_doc_to_docx(doc_path):
    word = Dispatch("Word.Application")
    
    doc = word.Documents.Open(doc_path)
    docx_path = os.path.splitext(doc_path)[0] + ".docx"
    
    doc.SaveAs(docx_path, FileFormat=16)
    
    doc.Close()
    word.Quit()

    return docx_path

def extract_text_from_pdf(pdf_path):
    with pdfplumber.open(pdf_path) as pdf:
        text = ""
        for page in pdf.pages:
            text += page.extract_text()
    return text

    
def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    
    # Extract text from headers
    for section in doc.sections:
        for header in section.header.paragraphs:
            text += header.text + "\n"
            

    # Extract text from footers
    for section in doc.sections:
        for footer in section.footer.paragraphs:
            text += footer.text + "\n"
            
    return text


def extract_email_addresses(text):
    email_regex = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    return re.findall(email_regex, text)

def extract_phone_numbers(text):
    phone_regex = r'(?:\+\d{1,2}\s?)?(?:\d{1,4}[-.\s]?)?\d{3,4}[-.\s]?\d{4}'


    return re.findall(phone_regex, text)

def extract_information(cv_path):
    if cv_path.endswith('.pdf'):
        text = extract_text_from_pdf(cv_path)
    elif cv_path.endswith('.doc'):
        text = extract_text_from_docx(convert_doc_to_docx(cv_path))
        os.remove(cv_path+'x')
    elif cv_path.endswith('.docx'):
        text = extract_text_from_docx(cv_path)
    else:
        raise ValueError("Unsupported file format")

    email_addresses = extract_email_addresses(text)
    phone_numbers = extract_phone_numbers(text)
    
    if len(phone_numbers) == 0 or len(email_addresses) == 0:
        print("Error at -->", cv_path)
        
    return {
        'Email': email_addresses[0],
        'Phone Number': phone_numbers[0],
        'Text': text
    }

def save_to_excel(data, excel_path):
    try:
        wb = load_workbook(excel_path)
        ws = wb.active
    except FileNotFoundError:
        wb = Workbook()
        ws = wb.active
        ws.append(['Email', 'Phone Number', 'Text'])
        
    for item in data:
        ws.append([item['Email'], item['Phone Number'], item['Text']])

    wb.save(excel_path)
    

if __name__ == "__main__":
    directory = input('Enter the directory: ')
    list_of_files = list_files_in_directory(directory)
    
    if list_of_files:
        cv_paths = [os.path.join(directory.replace('\\', '\\'), i) for i in list_of_files]
        n = len(cv_paths)
        
        print('files', cv_paths)
        
        output_excel_path = "output.xlsx"
        
        all_extracted_info = []
        for i in range(n):
            extracted_info = extract_information(cv_paths[i])
            all_extracted_info.append(extracted_info)
            print(f'Completed {i+1} out of {n}')

        save_to_excel(all_extracted_info, output_excel_path)
    else:
        print('No files in the directory')